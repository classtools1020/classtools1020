# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
style.font.size = Pt(11)

def set_cn(run, bold=False, size=11, color=None):
    run.font.name = 'Microsoft JhengHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),color)
    tcPr.append(sh)

def para(text, bold=False, size=11, color=None, align=None, space=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space)
    set_cn(p.add_run(text), bold, size, color)
    return p

def page_break():
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

# ====== 封面 ======
para('🤫 必須說出的祕密｜6/18 教案包', True, 20, RGBColor(0x6A,0x1B,0x9A), WD_ALIGN_PARAGRAPH.CENTER, 2)
para('國中特教班・社會科｜45分鐘｜好祕密 vs 壞祕密・勇敢說出來・找信任的大人', False, 11, RGBColor(0x55,0x55,0x55), WD_ALIGN_PARAGRAPH.CENTER, 4)
para('內含：① 上課逐字稿　② 學習單　③ 勇敢說出來求助卡（手作）', True, 11, RGBColor(0x1F,0x61,0x8C), WD_ALIGN_PARAGRAPH.CENTER, 10)

# ====== 逐字稿 ======
para('一、上課逐字稿', True, 16, RGBColor(0x6A,0x1B,0x9A))
t = doc.add_table(rows=1, cols=3); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['環節（時間）','🗣️ 你說（台詞）','🙆 你做/學生']):
    shade(t.rows[0].cells[i], '8E24AA')
    p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_cn(p.add_run(h), True, 11, RGBColor(0xFF,0xFF,0xFF))

rows=[
 ('① 暖身\n約3分',
  '「上課～老師問你們：你有沒有『祕密』呀？（等學生回答）每個人都有祕密對不對～但你知道嗎，祕密有兩種喔！」',
  '輕鬆聊天'),
 ('② 好祕密vs壞祕密\n約5分',
  '「🟢好祕密：讓人開心、是驚喜，像生日禮物、驚喜派對，之後會說出來。\n🔴壞祕密：讓你害怕、不舒服，還有人叫你『不能告訴任何人』——這種一定要說出來！」',
  '兩種顏色比手勢'),
 ('③ 危險訊號\n約3分',
  '「注意喔！如果有人說『這是我們的祕密，不可以告訴別人』，或用禮物、威脅叫你保密——這就是危險訊號！你的身體也會報警：睡不著、肚子痛、怕怕的。」',
  '提醒語氣'),
 ('④ 看影片\n約5分',
  '「我們看一段影片，記住一句話：『我說可以，才可以！』」（播放《我說可以才可以》易讀版）\n看完問：「誰可以決定？對～自己！」',
  '全班觀看\nlXCrxuEl-Kw'),
 ('⑤ 闖關判斷×6\n約12分',
  '節奏：唸題目→學生猜🟢好/🔴壞→翻頁揭曉→一句重點。\n①生日驚喜→🟢 ②有人碰你叫你別說→🔴 ③答應了讓你害怕的祕密→🔴（就算答應也要說）④祕密讓你睡不著肚子痛→🔴 ⑤網路叫你傳照片別說→🔴 ⑥威脅你別說出去→🔴',
  '每題翻頁揭曉\n答對：「太棒了👏」'),
 ('⑥ A組媒體識讀\n約5分',
  'A組：看《不能說的祕密》→ 識讀四問：①事實還是意見？②有沒有保護當事人？③標題會誇大嗎？④我能做什麼？\nB組同時：角色扮演「我要告訴老師」。',
  '分組進行'),
 ('⑦ 找信任的大人\n約4分',
  '「壞祕密要找『信任的大人』說：爸媽、老師、阿公阿嬤、輔導老師。說不出口怎麼辦？可以寫紙條，或打『113』——24小時都有人幫你！」',
  '帶全班念113'),
 ('⑧ 手作：求助卡\n約6分',
  '「現在做一張『勇敢說出來求助卡』，寫上你的3個信任大人和113，裝飾漂亮，帶在身邊。」',
  '發卡、巡桌、肯定'),
 ('⑨ 收尾金句\n約2分',
  '「記得：好祕密🟢可以收，壞祕密🔴一定說。說出來不是告密，是保護自己！」\n全班一起喊：「必須說出的祕密，我會勇敢說！」',
  '全班喊金句'),
]
for seg,say,do in rows:
    cells=t.add_row().cells
    for j,txt in enumerate([seg,say,do]):
        if j==0: shade(cells[j],'F3E5F5')
        cells[j].paragraphs[0].text=''
        first=True
        for line in txt.split('\n'):
            p=cells[j].paragraphs[0] if first else cells[j].add_paragraph(); first=False
            set_cn(p.add_run(line), bold=(j==0), size=10)
for i,w in enumerate([1.4,4.0,1.6]):
    for cell in t.columns[i].cells: cell.width=Inches(w)

# ====== 學習單 ======
page_break()
para('🤫 必須說出的祕密・學習單', True, 18, RGBColor(0x6A,0x1B,0x9A), WD_ALIGN_PARAGRAPH.CENTER, 4)
para('我的名字：________________', True, 13)

para('任務一：這是好祕密🟢還是壞祕密🔴？（圈圈看）', True, 13, RGBColor(0x1F,0x61,0x8C))
q1 = ['1. 幫朋友準備生日驚喜，先別說 …………… 🟢好 / 🔴壞',
      '2. 有人碰你的身體，叫你不能告訴爸媽 …… 🟢好 / 🔴壞',
      '3. 朋友送你的禮物先保密 ………………… 🟢好 / 🔴壞',
      '4. 讓你害怕、晚上睡不著的祕密 ………… 🟢好 / 🔴壞']
for q in q1: para(q, False, 12, space=3)

para('任務二：我的 3 個「信任的大人」（畫或寫）', True, 13, RGBColor(0x1F,0x61,0x8C))
tb = doc.add_table(rows=1, cols=3); tb.style='Table Grid'
for i in range(3):
    c=tb.rows[0].cells[i]; c.width=Inches(2.1)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_cn(p.add_run(f'\n\n（{i+1}）\n____________\n\n'), False, 12)

para('任務三：壞祕密三步驟（填 1、2、3）', True, 13, RGBColor(0x1F,0x61,0x8C))
para('（　）說出來　　（　）認出來，這讓我不舒服　　（　）找信任的大人', False, 12)

para('任務四：把金句寫一次', True, 13, RGBColor(0x1F,0x61,0x8C))
para('「說出來不是告密，是＿＿＿＿自己！」　　緊急電話：＿＿＿', False, 12)
para('⭐ 完成蓋章：★ ★ ★', True, 12, RGBColor(0xC0,0x39,0x2B), WD_ALIGN_PARAGRAPH.CENTER)

# ====== 求助卡手作 ======
page_break()
para('✋ 勇敢說出來・我的求助卡（手作）', True, 18, RGBColor(0x6A,0x1B,0x9A), WD_ALIGN_PARAGRAPH.CENTER, 4)
para('做法：寫上你的信任大人和電話，裝飾漂亮 → 護貝 → 帶在身邊！', False, 11, RGBColor(0x55,0x55,0x55), WD_ALIGN_PARAGRAPH.CENTER, 10)

card = doc.add_table(rows=1, cols=1); card.style='Table Grid'
shade(card.rows[0].cells[0], 'F3E5F5')
cc = card.rows[0].cells[0]
def cline(text, bold=False, size=13, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cc.add_paragraph(); p.alignment=align; p.paragraph_format.space_after=Pt(8)
    set_cn(p.add_run(text), bold, size, color)
cc.paragraphs[0].text=''
set_cn(cc.paragraphs[0].add_run('🆘 我的勇敢求助卡'), True, 16, RGBColor(0xC0,0x39,0x2B))
cc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
cline('我信任的大人：', True, 13)
cline('① ________________________', False, 13)
cline('② ________________________', False, 13)
cline('③ ________________________', False, 13)
cline('📞 緊急保護專線：113（24小時）', True, 14, RGBColor(0x1F,0x61,0x8C))
cline('💬 壞祕密，我會勇敢說出來！', True, 14, RGBColor(0x2E,0x7D,0x32), WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
para('🛠️ 材料：圖畫紙或厚卡紙、色筆、貼紙、護貝膜（或透明膠帶）。手部受限的學生可改用貼紙、蓋章代替寫字。', False, 11, RGBColor(0x55,0x55,0x55))

out='/home/user/classtools1020/0618必須說出的祕密_教案包.docx'
doc.save(out); print('saved', out)
