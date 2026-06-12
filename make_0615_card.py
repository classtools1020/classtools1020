# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
style.font.size = Pt(12)

def set_cn(run, bold=False, size=12, color=None):
    run.font.name = 'Microsoft JhengHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),color)
    tcPr.append(sh)

def para(text, bold=False, size=12, color=None, align=None, space=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space)
    set_cn(p.add_run(text), bold, size, color)
    return p

para('🌟 6/15 一對一教學・流程卡', True, 20, RGBColor(0xC0,0x39,0x2B), WD_ALIGN_PARAGRAPH.CENTER, 2)
para('主題：身體界線・感受紅綠燈｜對象：中度癲癇・手部功能受限・喜歡手作', False, 11, RGBColor(0x55,0x55,0x55), WD_ALIGN_PARAGRAPH.CENTER, 8)

# 安全提醒框
st = doc.add_table(rows=1, cols=1); st.style='Table Grid'
shade(st.rows[0].cells[0], 'FDECEA')
c = st.rows[0].cells[0].paragraphs[0]
set_cn(c.add_run('⚠️ 癲癇安全提醒\n'), True, 13, RGBColor(0xC0,0x39,0x2B))
set_cn(c.add_run('• 避免快閃／強閃光影片，以靜態圖與手作為主\n• 不用刀、不用火（材料免切免烤）\n• 放慢節奏；他嘆氣／低反應多為藥物嗜睡或疲累，非針對你\n• 事先確認發作徵兆與處理流程（IEP）；全程在旁陪同'), False, 11)
doc.add_paragraph()

# 心法
para('🫶 心法：一對一不是上課，是「我們一起做」。給他選擇、放慢、放大每個小成功。', True, 12, RGBColor(0x1F,0x61,0x8C), None, 8)

# 流程表
t = doc.add_table(rows=1, cols=3); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['環節（時間）','怎麼做','你說（引導語）']):
    shade(t.rows[0].cells[i], '4CAF50')
    p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_cn(p.add_run(h), True, 11, RGBColor(0xFF,0xFF,0xFF))

rows=[
 ('① 暖場・給選擇\n約3分',
  '拿出材料，讓他選今天做哪些：搖搖樂奶昔、捏捏球、黏土情緒臉（做2樣就好，時間到為準）。\n選擇＝參與感。',
  '「今天有三個好玩的～搖搖樂🥤、捏捏球🎈、黏土🎨，你想先做哪一個？你決定！」'),
 ('② 主活動A\n搖搖樂奶昔\n約15分',
  '材料：牛奶、香蕉(或草莓)、一點蜂蜜、搖搖杯。\n步驟：① 你幫他把材料放進杯 ② 蓋緊 ③ 他用力搖搖搖 ④ 倒出來喝。\n（動作＝大肌肉，手部受限也能做）',
  '「我們把材料放進去，蓋緊囉～現在用力搖！搖出好心情！」\n（喝的時候）「你今天是什麼顏色的心情呀？綠燈、黃燈、還是紅燈？」'),
 ('③ 主活動B\n減壓捏捏球\n約15分',
  '材料：氣球、麵粉或米、漏斗。\n步驟：① 用漏斗把麵粉灌進氣球(你協助) ② 打結 ③ 完成捏捏球，讓他盡情捏。\n＝帶得走的冷靜工具。',
  '「這是你的祕密武器～心裡紅燈、不舒服的時候，就捏一捏、深呼吸，等它慢慢變綠燈。我們一起捏捏看！」'),
 ('③＋ 主活動C\n黏土情緒臉\n（三選一備案）約15分',
  '材料：超輕黏土（紅黃綠三色，文具王購入，挑無毒、軟的）。\n玩法：① 捏一個「今天的心情」表情臉 ② 紅黃綠各搓一顆球，請他指「現在身體是哪一顆」（不用說話也能表達）。\n（捏、壓、搓＝大動作，省力療癒）',
  '「我們來捏一個『今天的你』～開心、累累的、還是悶悶的都可以。」\n「這三顆是紅綠燈～你現在的身體是哪一顆？指給老師看。」\n「不舒服的時候，捏一捏，慢慢深呼吸。」'),
 ('④ 收尾・連結主題\n約5分',
  '用作品聊「感受紅綠燈」：身體舒服=綠、不確定=黃、不舒服=紅。拍照、給大大的肯定。',
  '「今天你做得好棒！記得：你的身體、你的感受，你說了算。不舒服的時候，捏捏球＋告訴老師，好不好？」'),
]
for seg,do,say in rows:
    cells=t.add_row().cells
    for j,txt in enumerate([seg,do,say]):
        if j==0: shade(cells[j],'E8F5E9')
        cells[j].paragraphs[0].text=''
        first=True
        for line in txt.split('\n'):
            p=cells[j].paragraphs[0] if first else cells[j].add_paragraph(); first=False
            set_cn(p.add_run(line), bold=(j==0), size=10)
for i,w in enumerate([1.4,2.7,2.7]):
    for cell in t.columns[i].cells: cell.width=Inches(w)

doc.add_paragraph()
# 採購
para('🛒 全聯採購清單', True, 14, RGBColor(0x1F,0x61,0x8C))
para('搖搖樂：牛奶、香蕉或草莓、蜂蜜、搖搖杯（或寶特瓶）', False, 11)
para('捏捏球：氣球、麵粉或米、漏斗（氣球、漏斗 → 十元商店／文具店）', False, 11)
para('黏土：超輕黏土 紅・黃・綠 三色（文具王，挑無毒、軟的；避免太硬的油黏土）', False, 11)
para('共用：紙杯、湯匙、濕紙巾　｜　⚠️ 先確認牛奶/水果過敏', False, 11)

doc.add_paragraph()
para('💪 他用「垂的」、慢慢做都沒關係。你在旁邊陪、幫忙開蓋、放大每個小成功——陪伴本身就是最好的課。', True, 12, RGBColor(0xC0,0x39,0x2B))

out='/home/user/classtools1020/0615一對一教學流程卡.docx'
doc.save(out); print('saved', out)
