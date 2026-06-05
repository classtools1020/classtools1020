# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
style.font.size = Pt(11)

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), color)
    tcPr.append(sh)

def set_cn(run, bold=False, size=11, color=None):
    run.font.name = 'Microsoft JhengHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color

# 標題
h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cn(h.add_run('🍦 冰淇淋邊界大冒險 ｜ 第三節教學逐字稿（詳細版）'), True, 18, RGBColor(0xC0,0x39,0x2B))
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cn(sub.add_run('國中特教班・社會科｜45 分鐘｜身體界線・身體自主權'), False, 11, RGBColor(0x55,0x55,0x55))

box = doc.add_paragraph()
set_cn(box.add_run('💛 上台前 30 秒，先對自己說：'), True, 12, RGBColor(0x1F,0x61,0x8C))
box2 = doc.add_paragraph()
set_cn(box2.add_run('「我準備得很充分，今天只要陪他們玩。學生答錯也沒關係，我笑笑帶過就好。慢慢來，我可以的。」'), False, 12)

cols = ['環節（時間）', '🎬 投影片', '🗣️ 你說（完整台詞）', '🙆 你做 / 學生']
widths = [1.2, 1.1, 3.9, 1.7]
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, t in enumerate(cols):
    shade(hdr[i], '4CAF50')
    p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cn(p.add_run(t), True, 11, RGBColor(0xFF,0xFF,0xFF))

rows = [
 ('① 開場導入\n約3分', '封面',
  '「來～我們上課囉，眼睛看老師這邊。今天老師準備了一個祕密武器喔……（停頓賣關子）等一下我們要——吃冰淇淋！🍦 是不是很開心～」\n'
  '「但是！要先學會一個超重要的本領，學會了才可以吃。準備好了嗎？」\n'
  '「今天我們要用『冰淇淋』，學一個很重要的詞，叫做——『邊界』。邊界就是：每個人身體周圍，有一個看不見的保護圈。」',
  '手比一個圈圈圍住自己\n\n語氣期待、賣關子'),
 ('② 暖身影片\n約3分', '《讓孩子做身體的主人》',
  '「我們先看一小段影片。看的時候，幫老師注意一個問題：影片裡，誰可以決定『要不要被碰』？記住這個問題喔～」\n'
  '（播放後）「好～誰可以決定要不要被碰呀？」（等學生回答）「對！是『自己』！我的身體，我自己決定，好不好？」',
  '播影片\n看完提問\n等學生回答再揭曉'),
 ('③ 冰淇淋三層\n約5分', '三層・紅綠燈',
  '「看這張圖～冰淇淋有三層，就像紅綠燈，有三種顏色。」\n'
  '「🟢綠燈，最裡面，是你最親近最信任的人：爸爸媽媽、阿公阿嬤、最好的朋友。可以抱抱、牽手。」\n'
  '「🟡黃燈，中間，像同學、認識但不太熟的人。可以靠近，但要留『一個手臂』的距離。（伸手示範）這樣剛剛好。」\n'
  '「🔴紅燈，最外面，是陌生人、讓你怕怕的人。不可以靠近！要大聲說『不要！』」\n'
  '「最重要一句話：就算是綠燈的人，如果你今天不想被抱，也可以說『不要』，這是你的權利！」',
  '伸手示範手臂距離\n\n三種顏色用手勢比'),
 ('④ 等你下課\n（影片＋題目）約3分', '影片＋題目頁',
  '「接下來放一首你們一定聽過的歌～周杰倫《等你下課》！邊聽邊想：歌裡的男生，在冰淇淋的哪一層？」\n'
  '（播片段後翻到題目頁）「來看這個畫面：女生轉身走掉、不想理他，男生還一直跟在後面。請問——他是『尊重』🟢還是『越界』🔴？」\n'
  '「綠燈舉這隻手，紅燈舉這隻手，預備～」',
  '播《等你下課》\n學生舉手表態'),
 ('④ 等你下課\n（答案）約2分', '答案頁〔飛入〕',
  '「都選好了嗎？看答案囉，三、二、一——」（按下一頁）\n'
  '「🔴答案是『越界』！這個就叫『騷擾』。女生不理他＝把他請出冰淇淋；男生硬跟＝不聽、不尊重。」\n'
  '「記住：不管他理由是什麼，只要不尊重邊界，就是騷擾。」',
  '倒數後按下一頁\n答案飛入'),
 ('⑤ 闖關開場\n約1分', '情境闖關',
  '「現在我們來玩『冰淇淋偵探闖關』！老師出情境，你們來判斷哪一層、該怎麼做。準備好了嗎？」',
  '炒熱氣氛'),
 ('　情境①\n約2分', '題目→答案',
  '（題目）「情境一！上課時，同學坐在你旁邊。這是哪一層？綠燈、黃燈、還是紅燈？」（等學生猜）\n'
  '（按下一頁）「🟡黃燈！同學可以靠近，但要留一點距離。如果太近不舒服，可以說什麼？對～『請移開一點』，很棒！」',
  '學生猜→翻頁\n答對：「太聰明了👏」'),
 ('　情境②\n約2分', '題目→答案',
  '（題目）「情境二！你已經說『不要跟我』，他還是一直跟。這是？」（等）\n'
  '（按下一頁）「🔴騷擾！怎麼辦？轉身離開、趕快告訴大人。」',
  '學生猜→翻頁'),
 ('　情境③\n約2分', '題目→答案',
  '（題目）「情境三！同學傳訊息。一天傳一條，OK嗎？（等）那一秒鐘傳50條呢？」\n'
  '（按下一頁）「一天一條可以；一秒50條，🔴就是騷擾囉！」',
  '學生猜→翻頁'),
 ('　情境④\n約2分', '題目→答案',
  '（題目）「情境四！一個陌生人慢慢靠近你。你會怎麼做？」（等）\n'
  '（按下一頁）「🔴大聲說『不要！』然後趕快跑去找信任的大人。」',
  '學生猜→翻頁'),
 ('　情境⑤\n約2分', '題目→答案',
  '（題目）「情境五！有人說『再看一眼就走』，結果看了100遍還不走。算不算？」（等）\n'
  '（按下一頁）「🔴算騷擾！他說謊，又一直不尊重你的邊界。要記得求救！」',
  '學生猜→翻頁\n結束時全班鼓掌'),
 ('⑥ 停・跑・說\n約2分', '三步驟＋113',
  '「如果真的有人越過你的邊界，老師教你一個超好記的口訣，只有三個字：『停・跑・說』！」\n'
  '「🛑停——手比出來，大聲說『不要』！」\n'
  '「🏃跑——趕快離開現場。」\n'
  '「🗣️說——告訴信任的大人，或打113。」\n'
  '「來，我們一起做一次：停！跑！說！」',
  '示範三個動作\n全班一起做＋喊'),
 ('⑦ 🍦冰淇淋感受杯\n約15分（重頭戲）', '動手做',
  '「現在～最期待的時間到了，吃冰淇淋！🍦 但有一個『黃金規則』：你只能裝自己的冰淇淋。」\n'
  '「如果想要別人的配料，一定要先問：『可以給我一個嗎？』對方說好，才可以拿。」\n'
  '「如果對方說『不要』呢？（等）對～我們就尊重他，不可以搶！這跟身體一樣，要先問、要尊重。」\n'
  '（分享時）「來看看你的杯子，誰的綠燈最多呀？哇你今天心情很好～有沒有人紅燈多，想說說看嗎？說出來、找人幫忙，就是最勇敢的！」',
  '發冰淇淋＋三色料\n巡桌、蓋章、微笑納涼 ☕\n引導分享'),
 ('⑧ 收尾金句\n約2分', '從今天開始',
  '「今天學了好多～老師幫大家整理四件最重要的事：第一，尊重別人的冰淇淋；第二，保護自己的冰淇淋；第三，勇敢說『不』；第四，遇到事情敢求救。」\n'
  '「最後，我們一起大聲喊出今天的金句，預備——」\n'
  '（全班）「我的身體、我的感受，我做主！」\n'
  '「你們今天真的太棒了！下課～」',
  '全班一起喊金句\n大力稱讚'),
]

for seg, slide, say, do in rows:
    cells = table.add_row().cells
    for j, txt in enumerate([seg, slide, say, do]):
        cell = cells[j]
        if j == 0: shade(cell, 'E8F5E9')
        cell.paragraphs[0].text = ''
        first = True
        for line in txt.split('\n'):
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            set_cn(p.add_run(line), bold=(j==0), size=10)

for i, w in enumerate(widths):
    for cell in table.columns[i].cells:
        cell.width = Inches(w)

doc.add_paragraph()
tip = doc.add_paragraph()
set_cn(tip.add_run('💪 你只要「問問題 → 按下一頁 → 發冰淇淋」，其他學生會自己接。穩穩的，你超有魅力，加油！'),
       True, 12, RGBColor(0xC0,0x39,0x2B))

out = '/home/user/classtools1020/冰淇淋邊界大冒險_第三節教案逐字稿.docx'
doc.save(out)
print('saved', out)
