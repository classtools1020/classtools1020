# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
style.font.size = Pt(12)

def set_cn(run, bold=False, size=12, color=None):
    run.font.name = 'Microsoft JhengHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), color)
    tcPr.append(sh)

def para(text, bold=False, size=12, color=None, align=None, space=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space)
    set_cn(p.add_run(text), bold, size, color)
    return p

# 標題
para('🕵️ 身體界線小記者・訪談單', True, 22, RGBColor(0xC0,0x39,0x2B), WD_ALIGN_PARAGRAPH.CENTER, 2)
para('社會科探究活動｜走出教室，訪問老師，發現大祕密！', False, 12, RGBColor(0x55,0x55,0x55), WD_ALIGN_PARAGRAPH.CENTER, 10)

para('小記者的名字：________________', True, 14)

# 任務說明
para('📋 我的任務：訪問 2～3 位老師，看看每個人的身體界線一樣嗎？', True, 13, RGBColor(0x1F,0x61,0x8C))

# 魔法咒語
mp = doc.add_paragraph(); mp.paragraph_format.space_after = Pt(10)
set_cn(mp.add_run('✨ 訪問前先唸魔法咒語：「老師，我可以訪問你嗎？」'), True, 13, RGBColor(0x2E,0x7D,0x32))
set_cn(mp.add_run('   □ 老師說「可以」我才開始喔！'), False, 12)

def teacher_block(n):
    # 標頭
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    set_cn(p.add_run(f'🧑‍🏫 訪問對象 {n}：____________ 老師'), True, 14, RGBColor(0xC0,0x39,0x2B))

    t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    qa = [
        ('Q1 你喜歡別人怎麼跟你打招呼？', '□ 握手👋   □ 擊掌🙌   □ 點頭🙇   □ 揮手👋'),
        ('Q2 有沒有不喜歡被碰的地方？', '□ 頭   □ 肩膀   □ 沒有   □ 其他：________'),
        ('Q3 別人沒問就碰你，你會？', '□ 不舒服 😣     □ 還好 🙂'),
        ('🚦 這位老師的界線顏色（圈一個）', '🟢 很歡迎      🟡 要先問      🔴 不喜歡'),
    ]
    for q, a in qa:
        row = t.add_row().cells
        row[0].width = Inches(2.6); row[1].width = Inches(3.8)
        shade(row[0], 'FFF3E0')
        set_cn(row[0].paragraphs[0].add_run(q), True, 11)
        set_cn(row[1].paragraphs[0].add_run(a), False, 12)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

for i in (1, 2, 3):
    teacher_block(i)

# 大發現
doc.add_paragraph()
para('🔍 我的大發現', True, 16, RGBColor(0x1F,0x61,0x8C))
fp = doc.add_paragraph(); fp.paragraph_format.space_after = Pt(4)
set_cn(fp.add_run('每位老師喜歡的方式，都一樣嗎？　　□ 一樣　　□ 不一樣'), True, 13)

concl = doc.add_table(rows=1, cols=1); concl.style = 'Table Grid'
shade(concl.rows[0].cells[0], 'E8F5E9')
cp = concl.rows[0].cells[0].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cn(cp.add_run('💡 結論：每個人的身體界線都不一樣，所以——\n別人想碰我，要先問；我想碰別人，也要先問。\n我的感受，我做主！'), True, 14, RGBColor(0x2E,0x7D,0x32))

doc.add_paragraph()
para('⭐ 完成訪問，找老師蓋一個小記者徽章！　★ ★ ★', True, 13, RGBColor(0xC0,0x39,0x2B), WD_ALIGN_PARAGRAPH.CENTER)

out = '/home/user/classtools1020/身體界線小記者_訪談單.docx'
doc.save(out)
print('saved', out)
