# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
style.font.size = Pt(11)

def set_cn(run, bold=False, size=11, color=None):
    run.font.name='Microsoft JhengHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    run.font.size=Pt(size); run.font.bold=bold
    if color: run.font.color.rgb=color

def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),color); tcPr.append(sh)

def para(t,b=False,s=11,c=None,al=None,sp=6):
    p=doc.add_paragraph()
    if al: p.alignment=al
    p.paragraph_format.space_after=Pt(sp); set_cn(p.add_run(t),b,s,c); return p

def pb():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

PUR=RGBColor(0x6A,0x1B,0x9A)
para('🤫 必須說出的祕密｜教學逐字稿（教案版）', True, 20, PUR, WD_ALIGN_PARAGRAPH.CENTER, 2)
para('國中特教班・社會科｜好祕密 vs 壞祕密・身體界線・勇敢說出來', False, 11, RGBColor(0x55,0x55,0x55), WD_ALIGN_PARAGRAPH.CENTER, 4)
para('教學設計：引起動機 → 發展活動 → 練習 → 延伸 → 複習總結 → 手作', True, 11, RGBColor(0x1F,0x61,0x8C), WD_ALIGN_PARAGRAPH.CENTER, 8)

# 上台前
para('💛 上台前 30 秒：「我準備得很充分，今天只要陪他們玩。學生答錯也沒關係，我笑笑帶過。慢慢來，我可以的。」', True, 12, RGBColor(0x2E,0x7D,0x32))

def table(rows, headcolor='8E24AA'):
    t=doc.add_table(rows=1,cols=3); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(['教學階段（時間）','🗣️ 老師逐字稿','🙆 學生/做什麼']):
        shade(t.rows[0].cells[i],headcolor)
        p=t.rows[0].cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_cn(p.add_run(h),True,11,RGBColor(0xFF,0xFF,0xFF))
    for seg,say,do in rows:
        cells=t.add_row().cells
        for j,txt in enumerate([seg,say,do]):
            if j==0: shade(cells[j],'F3E5F5')
            cells[j].paragraphs[0].text=''
            first=True
            for line in txt.split('\n'):
                p=cells[j].paragraphs[0] if first else cells[j].add_paragraph(); first=False
                set_cn(p.add_run(line),bold=(j==0),size=10)
    for i,w in enumerate([1.5,4.0,1.6]):
        for cell in t.columns[i].cells: cell.width=Inches(w)
    return t

rows=[
 ('🎬 引起動機\n約5分',
  '（手藏一罐果醬在背後）「上課～老師今天帶了一個『祕密武器』，等一下要請你們吃東西喔！但先賣個關子～」\n'
  '「先問你們一個問題：你有沒有『祕密』呀？（等學生回答）每個人都有祕密對不對～'
  '但你知道嗎，祕密其實有『兩種』喔！今天我們就要當『祕密偵探』，把它們找出來！」',
  '老師賣關子、製造好奇\n學生回答有沒有祕密'),
 ('📖 發展活動一\n概念\n約8分',
  '「祕密有兩種，像紅綠燈：\n🟢好祕密：讓人開心，是驚喜（生日禮物、驚喜派對），之後會說出來。\n'
  '🔴壞祕密：讓你害怕、不舒服，還有人叫你『不能告訴任何人』——這種一定要說出來！」\n'
  '「要小心三個危險訊號：①有人說『這是我們的祕密不能說』 ②用禮物哄你保密 ③用威脅叫你保密。」\n'
  '「身體也會報警：睡不著、肚子痛、怕怕的——那就是身體在說『不對勁』！」',
  '兩色手勢🟢🔴\n師問：哪些是好/壞？'),
 ('🎥 發展活動二\n影片\n約5分',
  '「我們看一段影片，記住一句話：『我說可以，才可以！』看的時候想：誰可以決定要不要被碰？」\n'
  '（播《我說可以才可以》易讀版）看完問：「誰可以決定？對～自己！我的身體我做主！」',
  '全班看影片\n看完回答提問'),
 ('✋ 練習活動\n闖關判斷×6\n約12分',
  '節奏：唸題目 → 學生猜🟢好/🔴壞 → 翻頁揭曉 → 一句重點。\n'
  '①生日驚喜先別說→🟢　②有人碰你叫你別說→🔴　③答應了讓你害怕的祕密→🔴（就算答應也要說）\n'
  '④祕密讓你睡不著肚子痛→🔴　⑤網路叫你傳照片別說→🔴　⑥威脅你別說出去→🔴\n'
  '每題答對：「太聰明了，給自己拍拍手 👏」',
  '舉牌/手勢搶答\n翻頁揭曉驚喜'),
 ('🔍 延伸活動\n（A組）+ 找幫手\n約6分',
  'A組：看《不能說的祕密》→ 媒體識讀四問：①事實還是意見？②有沒有保護當事人？③標題誇大嗎？④我能做什麼？\n'
  '全班：「壞祕密要找『信任的大人』說：爸媽、老師、阿公阿嬤、輔導老師。說不出口可以打『113』——24小時都有人幫你！」',
  'A組分組討論\nB組角色扮演說出來'),
 ('🔁 複習與總結\n約4分',
  '「今天記住三步驟：認出來 → 說出來 → 找大人。」\n'
  '「還有最重要的：好祕密🟢可以收，壞祕密🔴一定說。說出來不是告密，是保護自己！」\n'
  '（全班一起喊）「必須說出的祕密，我會勇敢說！」',
  '全班一起做三步驟\n大聲喊金句'),
]
para('一、教學流程逐字稿', True, 15, PUR)
table(rows)

# 手作
pb()
para('二、手作活動逐字稿：身體吐司 🍞', True, 16, PUR)
para('核心比喻：吐司＝我的身體，果醬＝別人的接觸。一份吐司，吃進「身體自主權＋同意＋好壞祕密」。', True, 12, RGBColor(0x1F,0x61,0x8C))

mrows=[
 ('① 發吐司\n約2分',
  '「現在發給每個人一片吐司。注意喔——這片吐司，現在是『你的身體』。你決定上面要加什麼，因為『你的身體你做主』！」',
  '一人一片吐司、一支壓舌板'),
 ('② 先問規則\n（同意）約3分',
  '「想幫旁邊同學加果醬嗎？不可以直接加喔！要先問：『可以幫你加草莓嗎？』他說好才行。\n'
  '如果他說『不要』呢？（等）對～我們就尊重他，不可以硬加！這就跟身體一樣，別人要碰我，要先問。」',
  '兩兩練習「先問」\n示範說好/說不要'),
 ('③ 抹出心情\n約5分',
  '「🟢綠果醬＝我喜歡、舒服的；🔴紅果醬＝我不要的。自己抹出今天的心情吐司。」\n'
  '（巡桌）「哇你綠色好多，今天心情不錯喔～」',
  '自己抹、選顏色\n老師巡桌、肯定'),
 ('④ 串到壞祕密\n約3分',
  '「如果有人『硬在你的吐司上加你不要的東西』，還說『不能告訴老師喔』——這就是🔴壞祕密！\n'
  '你可以說『不要』，然後告訴信任的大人。記得嗎？認出來→說出來→找大人！」',
  '回應：要說出來！'),
 ('⑤ 驚喜吐司\n（好祕密）約3分',
  '「最後，我們偷偷做一份『驚喜吐司』送給好朋友～這種讓人開心、之後會公開的，就是🟢好祕密！」\n'
  '（享用）「謝謝大家今天當最棒的祕密偵探，開動囉！」',
  '做驚喜吐司送同學\n一起享用'),
]
table(mrows, headcolor='C0392B')

para('⚠️ 安全：吐司免烤最安全（要烤請老師操作）；事先確認過敏（奶、堅果、麩質），備一份替代。', False, 11, RGBColor(0xC0,0x39,0x2B))

# 準備清單
pb()
para('三、我要準備什麼？（準備清單）', True, 16, PUR)
para('🖥️ 教學設備', True, 13, RGBColor(0x1F,0x61,0x8C))
for x in ['投影設備、簡報（Google Slides）','影片先測試能播：《我說可以才可以》、《不能說的祕密》']:
    para('☐ '+x, False, 12, sp=3)
para('🎴 教學用具', True, 13, RGBColor(0x1F,0x61,0x8C))
for x in ['🟢🔴 舉牌或手勢卡（闖關用）','「113」提示卡、信任大人海報','（可選）五指安全網手掌學習單']:
    para('☐ '+x, False, 12, sp=3)
para('🍞 身體吐司材料（依人數）', True, 13, RGBColor(0x1F,0x61,0x8C))
for x in ['白吐司（免烤）','🔴 草莓果醬　🟡 柑橘/鳳梨果醬　🟢 奇異果/蘋果果醬','壓舌板或塑膠刀（抹醬，安全不傷手）','紙盤、濕紙巾','★ 先調查過敏：奶、堅果、麩質，備替代品']:
    para('☐ '+x, False, 12, sp=3)

para('', False, 6)
para('💪 你只要「賣關子 → 問問題 → 翻頁揭曉 → 發吐司」，學生會自己接。穩穩的，你超有魅力！', True, 12, RGBColor(0xC0,0x39,0x2B))

out='/home/user/classtools1020/0618必須說出的祕密_逐字稿教案版.docx'
doc.save(out)
import shutil; shutil.copy(out,'/home/user/classtools1020/0618_secrets_full_lesson_script.docx')
print('saved', out)
